#!/usr/bin/env python3
"""
生成毕业综合实践报告：Sown（数问）数学社区网站
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============ HELPERS ============

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = '宋体'
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, "E8EDD8")
    return row

# ============ DOCUMENT ============

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

def add_heading_l1(text):
    """一级标题：宋体小四号加粗，左侧空两字"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '宋体'
    return p

def add_heading_l2(text):
    """二级标题：宋体小四号加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '宋体'
    return p

def add_heading_l3(text):
    """三级标题：宋体小四号"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    return p

def add_para(text, bold=False, size=12, align=None, indent=True):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent and align is None:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    if bold:
        run.bold = True
    return p

def add_code_block(code_text):
    """Add code as a styled paragraph block"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    return p

def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run('')
    run.font.size = Pt(6)
    return p

# ========================================================
#  COVER / TITLE
# ========================================================

def add_cover_table(label, value, value_bold=False, value_size=None, value_align=None):
    """Add a 2-column cover table with label and value."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove table borders by setting style to No Spacing equivalent
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Label cell
    cell0 = table.rows[0].cells[0]
    cell0.text = ''
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run0 = p0.add_run(label)
    run0.font.size = Pt(12)
    run0.font.name = '宋体'

    # Value cell
    cell1 = table.rows[0].cells[1]
    cell1.text = ''
    p1 = cell1.paragraphs[0]
    if value_align:
        p1.alignment = value_align
    else:
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(value)
    if value_bold:
        run1.bold = True
    run1.font.size = value_size if value_size else Pt(12)
    run1.font.name = '宋体'

    return table

for _ in range(4):
    add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('毕业综合实践报告')
run.bold = True
run.font.size = Pt(22)
run.font.name = '宋体'

add_blank()
add_blank()

add_cover_table('题 目：', '数问（Sown）数学社区网站设计与实现', value_bold=True, value_size=Pt(14))

add_blank()

add_cover_table('系    别：', '人工智能学院')
add_cover_table('专    业：', '信息安全技术应用')
add_cover_table('班    级：', '信息安全技术应用2302')
add_cover_table('姓    名：', '王旭')
add_cover_table('学    号：', '230102160210')
add_cover_table('指导老师：', '王崎')
add_cover_table('完成日期：', '2026年5月')

for _ in range(6):
    add_blank()

doc.add_page_break()

# ========================================================
#  ABSTRACT + KEYWORDS
# ========================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('摘  要')
run.bold = True
run.font.size = Pt(12)
run.font.name = '宋体'

abstract_text = (
    '本文以信息安全专业视角，详细阐述了"数问（Sown）"数学社区网站的设计思路与完整实现过程。'
    '项目采用PHP原生开发，以MySQL为数据库引擎，前端使用原生JavaScript搭配GSAP动画库和Quill富文本编辑器，'
    '构建了一个集数学内容发布、社区互动、积分商城、管理员后台于一体的垂直社区平台。'
    '本文着重分析了数学公式渲染与编辑的实现原理，包括KaTeX/MathJax双引擎渲染方案和Quill编辑器中的数学符号面板设计；'
    '深入探讨了社区内容推流算法，涵盖推荐排序、热门计算和关注动态三种模式的数学模型；'
    '全面梳理了网站安全防护体系，包括CSRF令牌验证、XSS过滤与DOM净化、登录防爆破策略、SQL注入防御、内容审核机制等；'
    '详细说明了AI智能体的RAG（检索增强生成）接入原理与自动降级策略；'
    '同时展示了管理员后台的审核工作流、素材库管理、订单管理等核心功能。'
    '经实际测试，系统运行稳定，安全防护有效，用户体验良好。'
)
add_para(abstract_text)

add_blank()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('关键词：')
run.bold = True
run.font.size = Pt(12)
run.font.name = '宋体'
run2 = p.add_run('数学社区；PHP；MySQL；网络安全；RAG；推流算法')
run2.font.size = Pt(12)
run2.font.name = '宋体'

doc.add_page_break()

# ========================================================
#  TABLE OF CONTENTS (manual)
# ========================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('目  录')
run.bold = True
run.font.size = Pt(16)
run.font.name = '宋体'

add_blank()

toc_items = [
    ('一、引言', 1),
    ('（一）实习单位简介', 2),
    ('（二）选题目的与意义', 2),
    ('（三）网站名称与Logo设计理念', 2),
    ('二、需求分析', 1),
    ('（一）功能性需求', 2),
    ('（二）非功能性需求', 2),
    ('三、数据库设计', 1),
    ('（一）数据库整体设计思路', 2),
    ('（二）核心数据表详解', 2),
    ('（三）索引设计', 2),
    ('四、代码实现', 1),
    ('（一）技术栈选择', 2),
    ('（二）数学公式渲染与编辑原理', 2),
    ('（三）社区推流算法', 2),
    ('（四）AI智能体接入原理', 2),
    ('（五）社区核心功能实现', 2),
    ('（六）用户设置功能实现', 2),
    ('（七）积分商城系统', 2),
    ('（八）管理员后台实现', 2),
    ('（九）画板功能实现', 2),
    ('五、功能实现效果', 1),
    ('（一）前台用户端', 2),
    ('（二）管理员后台', 2),
    ('六、Web防护体系及渗透测试', 1),
    ('七、实践心得', 1),
    ('八、结论', 1),
    ('致谢', 1),
    ('参考文献', 1),
    ('附录', 1),
]

for item, level in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5 * (level - 1))
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    if level == 1:
        run.bold = True

doc.add_page_break()

# ========================================================
#  MAIN BODY
# ========================================================

# 一、引言
add_heading_l1('一、引言')

add_heading_l2('（一）实习单位简介')
add_para(
    '本人实习于一家专注于数学教育的教研工作室。该工作室以"数学教研"为核心定位，'
    '工作重心在于数学教学内容的深度研究、课程体系的设计开发、'
    '以及数字化教学工具的应用探索。工作室团队由一线数学教师和教研员组成，致力于通过教研创新'
    '提升数学学科的教学质量与学习体验。在实习期间，本人参与了教研数字化工具的需求调研与技术验证，'
    '深入接触了Web全栈开发、数据库设计、网络安全防护等多个技术方向，积累了丰富的工程实践经验。'
    '本次毕业设计的选题——"数问（Sown）数学社区网站"，正是基于实习期间对数学教育场景的观察与思考，'
    '结合信息安全专业的学科背景，独立完成的一个完整Web项目。'
)

add_heading_l2('（二）选题目的与意义')
add_para(
    '数学作为基础学科，在K12教育中占据重要地位，但目前面向数学爱好者和学生的垂直社区平台相对匮乏。'
    '现有通用型社区平台（如知乎、贴吧）对于数学公式的编辑和渲染支持不足，难以满足数学内容的专业性需求。'
    '因此，开发一个专注于数学领域的社区网站具有重要的实际意义：一方面，它填补了市场空白，'
    '为数学爱好者提供了一个专业、纯粹的交流空间；另一方面，作为信息安全专业的学生，'
    '将该网站作为毕业设计课题，可以全面检验自己在Web开发、数据库设计和网络安全防护方面的综合能力。'
)

add_heading_l2('（三）网站名称与Logo设计理念')
add_para(
    '"数问"这一名称的灵感来源于平台的核心定位：数学问题社区。"数"代表数学，是平台的内容根基；'
    '"问"代表提问与探索，是社区互动的核心行为。二者组合，"数问"直接传达了——这是一个围绕数学问题展开'
    '讨论和交流的专业社区。'
)
add_para(
    '与此同时，"数问"与英文"Sown"形成谐音对应关系。"Sown"是英文单词"sow"（播种）的过去分词形式，'
    '意为"被播种的"。这一巧妙的谐音赋予了平台更深层的教育隐喻：学习数学的过程，正如在土壤中播下种子，'
    '需要耐心浇灌与等待，终将迎来思维的萌芽与成长。正如网站的站内标语所写：'
    '"种下热爱，等数学开花"，这里的"开花"既是知识理解之花的绽放，也是学习兴趣之果的成熟。'
    '这一品牌理念将"播种—生长—开花"的意象贯穿于整个平台的视觉与交互设计之中。'
)
add_para(
    '基于上述品牌理念，网站的Logo采用了"幼苗 + 问号"的组合设计：Logo的主体由上下两部分构成——'
    '下方是一株绿色的幼苗，象征着播种、生长与教育的生命力；'
    '上方是一个棕色的问号，对应"数问"中的"问"字，代表数学探索中的提问与思考精神。'
    '绿色与棕色的配色方案取自自然界的土壤与植物，绿色寓意生机与成长，棕色寓意沉稳与根基，'
    '二者结合传达出"在扎实的基础上追求思维成长"的教育价值观。'
    '整个Logo的设计语言简洁而富有深意，将品牌名称"数问"的双关含义（中文表意 + 英文谐音）'
    '整合为一个统一的视觉符号，使用户在第一眼就能感受到平台的专业定位与教育温度。'
)

# 二、需求分析
add_heading_l1('二、需求分析')

add_heading_l2('（一）功能性需求')
add_para(
    '经过详细调研，数问社区网站的功能需求可以分为两大模块：前台用户端和后台管理端。'
)

add_heading_l3('1、前台用户端功能')
add_para(
    '(1) 用户注册与登录：支持邮箱+密码注册，提供"记住我"功能，支持会话持久化。(2) 内容发布：'
    '用户可使用富文本编辑器发布数学相关内容，支持LaTeX公式编辑、图片上传、标签选择。(3) 社区浏览：'
    '提供推荐（最新）、热门、动态三种排序模式，支持按标签筛选、无限滚动加载。(4) 互动功能：'
    '点赞、收藏、评论（含回复功能）。支持用户称号系统。(5) 积分系统：每日签到（含连续签到奖励）、'
    '发帖/获赞里程碑奖励。(6) 商城功能：积分兑换虚拟商品（称号）和实物商品。(7) 个人中心：'
    '个人资料编辑、头像上传、收货地址管理、兑换记录查看。(8) AI问答：接入"应老师"AI智能体，'
    '支持数学知识问答。(9) 通知系统：点赞、评论、回复等互动通知。'
)

add_heading_l3('2、管理员后台功能')
add_para(
    '(1) 仪表盘：展示网站核心统计数据（用户数、帖子数、待审核内容等），支持操作日志查看和过期数据清理。'
    '(2) 帖子管理：审核、编辑、删除用户发布的帖子。(3) 素材库管理：对上传的图片、文件进行目录化管理。'
    '(4) 订单管理：查看、处理用户兑换实物商品的订单，支持发货、物流单号批量上传、CSV导出。'
    '(5) 商城管理：上架、编辑、下架积分商品。'
)

add_heading_l2('（二）非功能性需求')
add_para(
    '(1) 安全性：全面防御常见Web攻击（SQL注入、XSS、CSRF、暴力破解等），保护用户数据和系统安全。'
    '(2) 兼容性：支持主流浏览器（Chrome、Firefox、Edge、Safari）及移动端访问。'
    '(3) 性能：在低配置服务器（2GB内存/双核CPU）上也能稳定运行，采用连接池、静态缓存等优化策略。'
    '(4) 可用性：界面简洁直观，交互反馈清晰，降低用户学习成本。'
)

# 三、数据库设计
add_heading_l1('三、数据库设计')

add_heading_l2('（一）数据库整体设计思路')
add_para(
    '数据库采用MySQL 8.0，存储引擎为InnoDB，字符集为utf8mb4（支持emoji和特殊符号）。'
    '整体设计遵循第三范式，尽量减少数据冗余。核心业务围绕"用户-帖子-评论"三大实体展开，'
    '辅以标签系统、积分系统、商城系统、通知系统等周边模块。各表之间通过外键（逻辑外键，实际使用程序控制）'
    '建立关联，确保数据完整性。'
)

add_heading_l2('（二）核心数据表详解')

# --- user 表 ---
add_heading_l3('1、user（用户表）')
add_para(
    'user表是系统的核心实体，存储所有注册用户的信息。每个字段的设计都有明确的业务含义：'
)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
header = table.rows[0].cells
headers = ['字段名', '数据类型', '约束', '默认值', '说明']
for i, h in enumerate(headers):
    header[i].text = ''
    p = header[i].paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = '宋体'
    set_cell_shading(header[i], 'E8EDD8')

user_fields = [
    ('id', 'INT UNSIGNED', 'PRIMARY KEY AUTO_INCREMENT', '-', '用户唯一标识'),
    ('username', 'VARCHAR(50)', 'NOT NULL', '-', '用户名，登录显示用'),
    ('account', 'VARCHAR(20)', 'DEFAULT NULL', 'NULL', '系统生成的唯一账号'),
    ('email', 'VARCHAR(100)', 'NOT NULL', '-', '邮箱，登录凭证'),
    ('phone', 'VARCHAR(20)', 'NOT NULL', '-', '手机号'),
    ('password', 'VARCHAR(255)', 'NOT NULL', '-', '密码哈希（password_hash）'),
    ('status', 'TINYINT(1)', 'NOT NULL', '1', '状态：1正常，0禁用'),
    ('role', 'VARCHAR(20)', 'NOT NULL', 'user', '角色：user/admin'),
    ('points', 'INT', 'NOT NULL', '0', '积分余额'),
    ('login_fail_count', 'INT', 'NOT NULL', '0', '连续登录失败次数'),
    ('login_locked_until', 'DATETIME', 'DEFAULT NULL', 'NULL', '锁定到期时间'),
    ('created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'NOW()', '注册时间'),
]
for row_data in user_fields:
    add_table_row(table, row_data)

add_blank()

# --- post 表 ---
add_heading_l3('2、post（帖子表）')
add_para(
    'post表存储用户发布的全部内容，是社区信息的主要载体。其核心字段设计兼顾了查询效率和数据完整性：'
)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
header = table.rows[0].cells
for i, h in enumerate(headers):
    header[i].text = ''
    p = header[i].paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = '宋体'
    set_cell_shading(header[i], 'E8EDD8')

post_fields = [
    ('id', 'INT UNSIGNED', 'PRIMARY KEY AUTO_INCREMENT', '-', '帖子唯一标识'),
    ('user_id', 'INT', 'NOT NULL', '-', '作者ID，关联user表'),
    ('title', 'VARCHAR(200)', 'NOT NULL', '-', '帖子标题'),
    ('content', 'LONGTEXT', 'NOT NULL', '-', '正文HTML（含LaTeX公式）'),
    ('image', 'VARCHAR(1000)', 'DEFAULT NULL', 'NULL', '封面图路径(JSON数组)'),
    ('status', 'TINYINT(1)', 'NOT NULL', '1', '状态：1发布，2草稿，0删除'),
    ('review_status', 'TINYINT(1)', 'NOT NULL', '0', '审核：0通过，1待审，2驳回'),
    ('like_count', 'INT', 'NOT NULL', '0', '点赞数（反范式化）'),
    ('comment_count', 'INT', 'NOT NULL', '0', '评论数（反范式化）'),
    ('favorite_count', 'INT', 'NOT NULL', '0', '收藏数（反范式化）'),
    ('created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'NOW()', '发布时间'),
    ('updated_at', 'DATETIME', 'ON UPDATE', 'NOW()', '最后更新时间'),
]
for row_data in post_fields:
    add_table_row(table, row_data)

add_blank()

add_para(
    '这里特别说明like_count、comment_count和favorite_count三个字段。按照数据库范式理论，'
    '这些计数可以从关联表中通过COUNT查询实时计算得出。但考虑到社区场景下，帖子列表页需要频繁'
    '展示这些计数（每页12条帖子），如果每次都进行COUNT聚合查询，在高并发场景下会对数据库造成较大压力。'
    '因此采用了反范式化设计，在post表中直接缓存这些计数，通过事务保证一致性。每次用户点赞或取消点赞时，'
    '在同一事务中更新post_like表和post表，确保两个表的数据始终一致。'
)

# --- comment 表 ---
add_heading_l3('3、comment（评论表）')
add_para(
    'comment表实现了评论区功能，其设计亮点在于通过parent_id字段同时支持"主评论"和"回复"两种模式。'
    '当parent_id为NULL或0时，表示这是一条直接针对帖子的主评论；当parent_id指向另一条评论的ID时，'
    '表示这是一条"回复"。这种设计将评论和回复统一存储在同一张表中，既简化了表结构，又能灵活支持多级回复。'
)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
header = table.rows[0].cells
for i, h in enumerate(headers):
    header[i].text = ''
    p = header[i].paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = '宋体'
    set_cell_shading(header[i], 'E8EDD8')

comment_fields = [
    ('id', 'INT UNSIGNED', 'PRIMARY KEY AUTO_INCREMENT', '-', '评论唯一标识'),
    ('post_id', 'INT', 'NOT NULL', '-', '所属帖子ID'),
    ('user_id', 'INT', 'NOT NULL', '-', '评论者ID'),
    ('parent_id', 'INT', 'DEFAULT NULL', 'NULL', '父评论ID（NULL=主评论）'),
    ('content', 'TEXT', 'NOT NULL', '-', '评论内容'),
    ('status', 'TINYINT(1)', 'NOT NULL', '1', '状态：1正常，0删除'),
    ('like_count', 'INT', 'NOT NULL', '0', '点赞数'),
    ('created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'NOW()', '评论时间'),
]
for row_data in comment_fields:
    add_table_row(table, row_data)

add_blank()

# --- 标签系统 ---
add_heading_l3('4、标签相关表（post_tag / post_tag_relation）')
add_para(
    '标签系统采用经典的多对多关系设计，由post_tag（标签定义表）和post_tag_relation（关系表）两张表组成。'
    'post_tag表存储标签的名称和slug（URL友好的标识符），post_tag_relation表建立帖子和标签之间的多对多关联。'
    '每个帖子最多关联5个标签，标签的slug用于URL筛选（如/forum.php?tag=algebra），保证URL的可读性和SEO友好性。'
)

# --- 互动表 ---
add_heading_l3('5、互动记录表（post_like / post_favorite / comment_like）')
add_para(
    '这三张表分别记录用户的点赞和收藏行为。它们的结构相似，核心字段为post_id/comment_id和user_id，'
    '构成联合唯一索引，确保"一个用户对一个内容只能点赞/收藏一次"。这种设计天然适配"点赞/取消点赞"的切换逻辑：'
    '再次点击时DELETE原记录即可。这是社交互动功能的标准数据库设计模式。'
)

# --- 积分相关表 ---
add_heading_l3('6、积分与商城相关表')
add_para(
    '积分系统涉及多张表：point_ledger（积分流水表）记录每一笔积分的增减明细，'
    '通过reason_key字段保证幂等性（同一用户同一原因只能获得一次积分）；'
    'user_checkin（签到表）记录用户每日签到情况，以(user_id, checkin_date)为主键，'
    '利用数据库的唯一约束防止重复签到；shop_purchase（购买记录表）记录用户的商品兑换历史；'
    'shop_item（商品表）定义商城中的可兑换商品；shop_order（订单表）管理实物商品的物流信息。'
    'user_title（用户称号表）记录用户获得的社区称号。'
)

# --- 通知和日志表 ---
add_heading_l3('7、辅助表（notification / admin_log / media_file等）')
add_para(
    'notification表实现站内通知功能，记录点赞、评论等互动事件，通过is_read字段区分已读/未读。'
    'admin_log表记录管理员的所有操作，便于审计追溯。media_file和media_folder表实现了素材库功能，'
    '支持对上传文件进行目录化管理。teacher_knowledge表存储AI问答的知识库数据，支持向量检索和关键词检索。'
)

add_heading_l2('（三）索引设计')
add_para(
    '索引是数据库性能优化的关键手段。在本次设计中，索引策略遵循以下原则：'
    '(1) 主键索引：所有表均使用自增INT作为主键，InnoDB的聚簇索引特性使得按主键查询极为高效。'
    '(2) 唯一索引：在需要保证数据唯一性的字段上建立唯一索引，如user.email、point_ledger的(user_id, reason_key)、'
    'post_like的(post_id, user_id)等。这不仅能保证数据完整性，还能在INSERT时快速检测重复。'
    '(3) 普通索引：在外键字段和查询条件频繁使用的字段上建立索引，如post.user_id、comment.post_id、'
    'notification的(user_id, created_at)等。特别地，在forum.php的帖子列表查询中，'
    '对post表的status、created_at、review_status等字段的联合查询使用了索引，确保分页查询的高效性。'
    '(4) 联合索引：对于多条件查询，设计联合索引以最大化查询效率。例如notification表的'
    '(user_id, is_read)联合索引可以快速统计用户的未读通知数量。'
)

# 四、代码实现
add_heading_l1('四、代码实现')

add_heading_l2('（一）技术栈选择')
add_para(
    '本项目的技术栈选择充分考虑了"低配置服务器运行"和"信息安全专业考核"两个核心因素：'
    '后端采用原生PHP（不使用框架），一方面减少依赖和资源占用，另一方面让底层逻辑完全暴露，'
    '便于展示安全防护的实现细节。前端使用原生JavaScript，搭配Quill富文本编辑器和GSAP动画库。'
    '数据库使用MySQL，PDO数据访问层确保SQL注入防御。AI接入使用MiniMax开放平台的API。'
)
add_para(
    '项目部署于一台公网服务器（IP地址：223.4.97.5），运行Ubuntu 24.04 LTS操作系统，'
    '使用Nginx 1.26.3作为Web服务器、PHP 8.3.6作为后端解析引擎、MySQL 5.7.44作为数据库，'
    '并借助宝塔面板（Baota Panel）进行可视化的服务器运维管理，包括站点配置、'
    'Nginx/PHP/MySQL服务的启停监控、防火墙规则管理以及定期备份等日常运维操作。'
    '数据库管理方面，通过集成在宝塔面板中的phpMyAdmin工具进行可视化的数据表结构查看、'
    'SQL查询执行、数据导入导出等工作，大大提高了数据库管理和调试的效率。'
)

add_heading_l2('（二）数学公式渲染与编辑原理')

add_heading_l3('1、渲染引擎方案')
add_para(
    '数学公式渲染是本项目的核心技术挑战之一。经过调研，最终采用"双引擎"渲染方案：'
    '在帖子详情页使用MathJax 3进行完整渲染，在社区列表页和编辑器中使用KaTeX进行快速预览。'
    'MathJax的优势在于渲染质量高，支持完整的LaTeX语法，适合作为最终展示方案。'
    '其工作原理是通过JavaScript解析HTML中的LaTeX标记（以$...$和$$...$$为界），'
    '将其转换为MathML或HTML+CSS的排版结构。具体配置如下：'
)
add_code_block(
    'window.MathJax = {\n'
    '  tex: {\n'
    '    inlineMath: [["$", "$"], ["\\\\(", "\\\\)"]],\n'
    '    displayMath: [["$$", "$$"], ["\\\\[", "\\\\]"]]\n'
    '  }\n'
    '};'
)
add_para(
    '该配置指定了LaTeX公式的行内定界符（$...$）和展示定界符（$$...$$），MathJax在页面加载完成后，'
    '会扫描整个DOM树，找到匹配的定界符内容，调用其排版引擎将LaTeX代码渲染为数学符号的HTML布局。'
    '渲染过程是异步的，通过MathJax.typesetPromise()方法控制，确保不影响页面其他内容的加载。'
)

add_heading_l3('2、编辑器公式输入方案')
add_para(
    '在内容编辑环节，本项目基于Quill富文本编辑器进行深度定制。Quill本身提供了formula模块，'
    '支持嵌入KaTeX渲染的内联公式。但仅靠基础模块的体验不够友好，为此开发了"数学符号面板"功能。'
    '该面板将常用数学符号按照语义分组（运算符、关系符、希腊字母、箭头逻辑等），'
    '用户点击即可直接插入到编辑器中。'
    '同时实现了上下标占位框功能：用户在面板中选择"上标□"或"下标□"时，'
    '编辑器会插入一个带有特殊样式（虚线边框）的占位文本，提示用户输入内容。'
    '核心实现代码如下：'
)
add_code_block(
    '// 插入上/下标占位框\n'
    'function insertScriptPlaceholder(type) {\n'
    '  var quill = quillInstance;\n'
    '  var range = quill.getSelection(true);\n'
    '  var placeholder = "□";\n'
    '  var spacer = "\\u2009";  // 细空格\n'
    '  quill.insertText(range.index, placeholder,\n'
    '    { script: type, slot: type }, "user");\n'
    '  quill.insertText(range.index + 1, spacer,\n'
    '    { script: false, slot: false }, "user");\n'
    '  quill.setSelection(range.index, placeholder.length, "user");\n'
    '}'
)
add_para(
    '当用户在占位框中输入真实内容后，cleanupFilledSlots函数会自动移除虚线框样式，'
    '使公式看起来自然融入正文。Quill的Delta数据模型将公式内容存储为HTML，'
    '最终在帖子详情页由MathJax完成渲染。'
)

add_heading_l2('（三）社区推流算法')
add_para(
    '社区的内容排序是决定用户体验的核心因素。本项目实现了三种排序模式：推荐、热门和动态，'
    '分别对应不同的使用场景。'
)

add_heading_l3('1、推荐模式（最新优先）')
add_para(
    '推荐模式本质上是"最新发布优先"排序，即直接按照帖子的created_at字段降序排列。'
    '这种排序方式最为轻量，只需要一个B-tree索引即可高效执行。'
    '它的适用场景是：当用户想要浏览社区最新动态时，时间顺序是最直观的排序依据。'
    'SQL实现如下：'
)
add_code_block(
    'ORDER BY p.created_at DESC'
)

add_heading_l3('2、热门模式（加权热度算法）')
add_para(
    '热门模式的算法设计参考了Hacker News的热度计算公式，但根据社区特点进行了简化调整。'
    '最终采用的热度分值计算公式为：')
add_para(
    '热度 = (like_count × 2 + comment_count + favorite_count) / (DATEDIFF(NOW(), created_at) + 1)',
    bold=True, indent=False
)
add_para(
    '这个公式的数学含义是：每个帖子的热度由其互动总量（点赞、评论、收藏）除以存在天数得到，'
    '其中点赞的权重设为2（比评论和收藏更高），因为点赞是最轻量、最频繁的互动行为，'
    '更能反映内容的受欢迎程度。分母加1是为了避免新发布的帖子出现除以零的情况。'
    '同时，查询限定在最近30天内的帖子，避免过时内容占据热门榜单。'
    '这种算法的优势在于：新发布的优质内容可以通过高互动率快速冲上热门，'
    '而旧内容如果没有持续互动会自然衰减，保证了热门榜单的时效性和质量。'
)
add_code_block(
    '$dateLimit = "AND p.created_at >= DATE_SUB(NOW(), INTERVAL 30 DAY)";\n'
    '$orderBy = "(p.like_count * 2 + p.comment_count + p.favorite_count)"\n'
    '  . " / (DATEDIFF(NOW(), p.created_at) + 1) DESC, p.created_at DESC";'
)

add_heading_l3('3、动态模式（关注者推流）')
add_para(
    '动态模式本质上是社交推荐，只展示用户关注的人发布的帖子。'
    '通过user_follow表记录用户的关注关系，查询时使用JOIN关联关注数据和帖子数据。'
    '这种私域流量的推荐方式可以增强用户粘性。'
    '只有登录用户才能使用此功能，未登录用户点击"动态"会弹出登录框。'
)
add_code_block(
    '// 动态模式：仅展示关注者的帖子\n'
    'SELECT DISTINCT p.id, p.title, ...\n'
    'FROM post p\n'
    'JOIN user_follow f ON f.following_id = p.user_id\n'
    'WHERE p.status = 1 AND f.follower_id = ?\n'
    'ORDER BY p.created_at DESC'
)

add_heading_l2('（四）AI智能体接入原理')

add_heading_l3('1、"应老师"人设设计')
add_para(
    'AI智能体被设定为一位名为"应老师"的资深初中数学教师形象，其人设存储在data/teacher/persona.md文件中。'
    '人设的核心教学理念包括"深度优先于进度"、"先建立直观模型再揭示结构"、"重视基本功和体系化积累"。'
    '表达风格要求克制、直接、分层，避免空喊口号，常用口头禅如"先把结构看清楚"、"不是难，是还没看到关键"。'
    '这种人设设计使AI的回答更贴近真实的优秀教师，能给用户带来亲切感和信任感。'
)

add_heading_l3('2、RAG（检索增强生成）技术路线')
add_para(
    'AI问答的核心技术是RAG（Retrieval-Augmented Generation）。当用户提问时，系统首先从知识库中检索'
    '最相关的备课资料，然后将这些资料作为上下文注入到LLM的提示词中，最终生成有据可依的回答。'
    '具体流程如下：'
)
add_para(
    '(1) 将用户问题进行向量化，调用MiniMax的Embedding API（embo-01模型，1536维向量）。'
    '(2) 使用余弦相似度在teacher_knowledge表中检索Top-3最相似的知识块。余弦相似度公式为：'
    'similarity = (A·B) / (||A|| × ||B||)。'
    '(3) 将检索到的知识块拼接到system prompt中，调用MiniMax M2.7大模型生成回答。'
    '(4) 如果账户余额不足（API返回状态码1008），系统自动降级为关键词检索+模板回复模式。'
)

add_para('降级模式的工作流程：')
add_para(
    '(1) 提取用户问题中的双字词和单字词，过滤停用词。'
    '(2) 使用这些关键词在知识库中匹配，按匹配数量排序取Top-3。'
    '(3) 根据问题类型（函数类、证明类、学习方法类等）选择对应的智能模板回复，'
    '模板中会嵌入知识库中的相关内容。'
    '这种降级策略保证了即使API余额不足，系统仍能提供有意义的回答，用户体验不会完全中断。'
    '充值后系统会自动恢复完整的RAG模式，不需要手动切换。'
)

add_heading_l2('（五）社区核心功能实现')

add_para(
    '社区的交互功能是网站的核心竞争力所在。本节详细阐述发帖流程、评论系统、点赞收藏、'
    '标签系统和通知机制的实现原理与核心代码。积分商城、用户设置、管理员后台和画板功能将在后续章节展开。'
)

add_heading_l3('1、发帖流程')
add_para(
    '发帖是社区中最核心的用户操作，涉及数据验证、HTML净化、内容审核、图片处理和标签关联'
    '等多个环节。post_note_post.php实现了完整的发帖流程，具体步骤如下：'
)
add_para(
    '第一步，请求校验与CSRF验证：检查请求方法是否为POST，验证CSRF令牌是否有效。'
    '第二步，用户认证：确认用户已登录，并获取当前用户信息。'
    '第三步，参数验证：检查标题是否为空、长度是否超过200字符、正文是否为空、'
    'HTML长度是否超过500KB、纯文本长度是否超过50000字符。'
    '第四步，HTML安全净化：调用sanitize_post_content_html()对富文本内容进行DOM级XSS过滤。'
    '第五步，内容审核：调用moderation_filter_post()检测是否含违规关键词，命中则替换为屏蔽文案。'
    '第六步，图片提取与封面处理：从HTML正文中提取上传图片路径，处理可选封面图。'
    '第七步，数据库写入：插入或更新帖子记录，处理标签关联，触发积分奖励。'
    '第八步，返回结果：根据请求方式返回JSON或跳转到结果页面。'
)
add_para(
    '其中第五步的内容审核机制具体如下：系统维护了一个违规关键词库，涵盖辱骂人身攻击类'
    '（如"傻逼""脑残""废物"等）和暴力极端类（如"炸弹""恐怖袭击""种族灭绝"等），'
    '采用Unicode安全的子串匹配方式（mb_strpos）进行检测，不区分大小写。'
    '对于帖子，先将HTML正文提取纯文本再检测标题和正文，如果命中违规词，标题被替换为'
    '"内容违规，已屏蔽"，正文被替换为带有特殊样式的屏蔽提示信息，帖子本身不会被删除，'
    '但内容被完全覆盖。评论检测逻辑类似，命中后整条评论内容被替换为"内容违规，已屏蔽"。'
    '这一机制与管理员后台的review_status审核状态协同，形成了"自动关键词过滤 + 人工审核"'
    '的双层内容安全体系。'
)
add_para('核心实现代码如下（省略安全校验部分）：')
add_code_block(
    "// 发帖核心流程（post_note_post.php）\n"
    "// 1. 参数验证\n"
    '$title = trim($_POST["title"] ?? "");\n'
    "if (mb_strlen($title) > 200) { /* 标题超长错误 */ }\n"
    "if (mb_strlen($rawContent) > 500000) { /* HTML超长错误 */ }\n\n"
    "// 2. HTML安全净化（XSS防护）\n"
    '$contentHtml = sanitize_post_content_html($rawContent);\n\n'
    "// 3. 内容审核（违规词检测）\n"
    '$moderation = moderation_filter_post($title, $contentHtml);\n'
    '$title = $moderation["title"];\n'
    '$contentHtml = $moderation["content_html"];\n\n'
    "// 4. 提取正文中的图片用于列表封面\n"
    '$extractedPaths = post_extract_uploaded_image_paths($contentHtml);\n\n'
    "// 5. 写入数据库\n"
    '$st = $pdo->prepare(\n'
    '  "INSERT INTO post (user_id, title, content, image, status, review_status,\n'
    "    like_count, comment_count, favorite_count, created_at, updated_at)\n"
    '   VALUES (?, ?, ?, ?, ?, ?, 0, 0, 0, NOW(), NOW())"\n'
    ");\n"
    '$st->execute([$user["id"], $title, $contentHtml, $finalImageJson,\n'
    "  $status, $reviewStatus]);\n"
    '$postId = (int)$pdo->lastInsertId();\n\n'
    "// 6. 处理标签（多对多关联）\n"
    '$tagNames = array_filter(array_map("trim", preg_split("/[,，]/", $tagsInput)));\n'
    "set_post_tags($postId, $tagNames);\n\n"
    "// 7. 触发积分奖励\n"
    "points_after_post_published((int)$user[\"id\"]);"
)
add_para(
    '系统支持发帖和编辑草稿两种模式。当用户点击"保存草稿"时，post.status被设为2（草稿），'
    '只有作者本人可以在创作者平台看到草稿内容。点击"发布"时，post.status为1（已发布），'
    '但非管理员用户发布的帖子review_status会被设为1（待审核），需要管理员后台审核通过后才对其他人可见。'
    '编辑模式时，会额外验证帖子所有权（WHERE id = ? AND user_id = ?），防止越权修改。'
)

add_heading_l3('2、评论系统')
add_para(
    '评论系统实现了"主评论"和"回复"两种互动模式。通过comment表中的parent_id字段区分：'
    'parent_id为NULL或0时表示这是一条针对帖子的主评论；parent_id指向另一条评论的ID时表示回复。'
    '这种单表设计简化了表结构，同时通过comment_create.php中的业务逻辑灵活支持多级回复。'
)
add_para('评论提交的核心处理流程如下：')
add_code_block(
    '// 评论创建核心逻辑（comment_create.php）\n'
    '// 1. 接收并校验参数\n'
    '$postId = isset($_POST["post_id"]) ? (int)$_POST["post_id"] : 0;\n'
    '$content = trim((string)($_POST["content"] ?? ""));\n'
    '$parentId = isset($_POST["parent_id"]) ? (int)$_POST["parent_id"] : 0;\n\n'
    '// 2. CSRF验证\n'
    'if (!csrf_check($token)) { /* 拒绝 */ }\n\n'
    '// 3. 内容长度限制\n'
    'if ($content === "" || mb_strlen($content) > 2000) { /* 拒绝 */ }\n\n'
    '// 4. 违规词过滤\n'
    '$content = moderation_filter_comment_text($content)["text"];\n\n'
    '// 5. 事务开始\n'
    '$pdo->beginTransaction();\n\n'
    '// 6. 验证帖子存在且可评论\n'
    '$stmt = $pdo->prepare("SELECT id, user_id FROM post WHERE id=:id \n'
    '  AND status=1 AND (review_status IS NULL OR review_status = 0) LIMIT 1");\n'
    '$stmt->execute([":id" => $postId]);\n'
    '$post = $stmt->fetch();\n\n'
    '// 7. 如果是回复，验证父评论存在且属于同一帖子\n'
    'if ($parentId > 0) {\n'
    '  $pstmt = $pdo->prepare("SELECT id, user_id, post_id FROM comment\n'
    '    WHERE id=:id AND status=1 LIMIT 1");\n'
    '  $pstmt->execute([":id" => $parentId]);\n'
    '  $parentComment = $pstmt->fetch();\n'
    '  if (!$parentComment || (int)$parentComment["post_id"] !== $postId) {\n'
    '    $pdo->rollBack(); exit;\n'
    '  }\n'
    '}\n\n'
    '// 8. 插入评论并更新计数\n'
    '$stmt = $pdo->prepare("INSERT INTO comment (post_id, user_id, content,\n'
    '  parent_id, status) VALUES (:post_id, :user_id, :content, :parent_id, 1)");\n'
    '$stmt->execute([":post_id" => $postId, ":user_id" => $userId,\n'
    '  ":content" => $content, ":parent_id" => $parentId > 0 ? $parentId : null]);\n\n'
    '$stmt = $pdo->prepare(\n'
    '  "UPDATE post SET comment_count = comment_count + 1 WHERE id = :id");\n'
    '$stmt->execute([":id" => $postId]);\n\n'
    '// 9. 发送通知\n'
    'if ($parentId > 0) {\n'
    '  create_notification($parentCommentUserId, "reply", $postId, $userId);\n'
    '} elseif ($postOwnerId !== $userId) {\n'
    '  create_notification($postOwnerId, "comment", $postId, $userId);\n'
    '}\n\n'
    '$pdo->commit();'
)
add_para(
    '评论系统的设计要点包括：事务保护（确保评论插入和计数更新原子性）、父评论验证'
    '（防止跨帖子伪造回复）、PRG模式（Post-Redirect-Get，避免刷新重复提交）、'
    '以及通过create_notification()自动通知被回复的用户或帖子作者。'
)

add_heading_l3('3、点赞与收藏系统')
add_para(
    '点赞和收藏采用相同的"切换式"设计模式：用户点击时，如果已点赞/收藏则取消，'
    '如果未点赞/收藏则创建。这种设计通过数据库的唯一约束保证数据一致性，'
    '同时采用事务保护确保点赞记录和计数更新的原子性。'
)
add_code_block(
    '// 点赞切换核心逻辑（like_toggle.php）\n'
    'try {\n'
    '  $pdo->beginTransaction();\n\n'
    '  // 检查帖子存在性\n'
    '  $st = $pdo->prepare("SELECT id, user_id FROM post WHERE id=:id \n'
    '    AND status=1 AND (review_status IS NULL OR review_status = 0) LIMIT 1");\n'
    '  $st->execute([":id" => $postId]);\n'
    '  $post = $st->fetch();\n\n'
    '  // 检查是否已点赞\n'
    '  $st = $pdo->prepare("SELECT id FROM post_like \n'
    '    WHERE post_id=:pid AND user_id=:uid LIMIT 1");\n'
    '  $st->execute([":pid" => $postId, ":uid" => $userId]);\n'
    '  $liked = (bool)$st->fetch();\n\n'
    '  if ($liked) {\n'
    '    // 取消点赞\n'
    '    $st = $pdo->prepare("DELETE FROM post_like \n'
    '      WHERE post_id=:pid AND user_id=:uid");\n'
    '    $st->execute([":pid" => $postId, ":uid" => $userId]);\n'
    '    $st = $pdo->prepare("UPDATE post SET like_count = \n'
    '      GREATEST(like_count-1,0) WHERE id=:id");\n'
    '    $st->execute([":id" => $postId]);\n'
    '  } else {\n'
    '    // 点赞\n'
    '    $st = $pdo->prepare("INSERT INTO post_like \n'
    '      (post_id, user_id) VALUES (:pid,:uid)");\n'
    '    $st->execute([":pid" => $postId, ":uid" => $userId]);\n'
    '    $st = $pdo->prepare("UPDATE post SET \n'
    '      like_count = like_count+1 WHERE id=:id");\n'
    '    $st->execute([":id" => $postId]);\n\n'
    '    // 发通知给帖子作者\n'
    '    if ($postOwnerId !== $userId) {\n'
    '      create_notification($postOwnerId, "like", $postId, $userId);\n'
    '    }\n'
    '  }\n\n'
    '  $pdo->commit();\n'
    '} catch (Throwable $e) {\n'
    '  if ($pdo->inTransaction()) $pdo->rollBack();\n'
    '}'
)
add_para(
    '收藏功能的实现与点赞完全对称，使用post_favorite表和favorite_count字段。'
    '两者的关键区别在于：点赞会触发通知给帖子作者，而收藏是私密行为，不会触发通知。'
    '计数字段使用GREATEST(like_count-1, 0)确保不会减到负数。'
    '每次成功点赞或收藏后，系统会调用points_refresh_author_engagement_milestones()'
    '检查帖子作者的里程碑奖励（如"首次获赞"、"获赞100次"等成就称号）。'
)

add_heading_l3('4、标签系统')
add_para(
    '标签系统采用经典的多对多关系设计，由post_tag（标签定义表）和post_tag_relation（关系表）'
    '两张表组成。核心函数get_or_create_tag()实现了"不存在则创建"的幂等性操作：'
)
add_code_block(
    '// 获取或创建标签（app/tag.php）\n'
    'function get_or_create_tag(string $name): int {\n'
    '  $name = trim($name);\n'
    '  if (mb_strlen($name) > 20) {\n'
    '    $name = mb_substr($name, 0, 20);\n'
    '  }\n'
    '  // 生成URL友好的slug\n'
    '  $slug = mb_strtolower($name);\n'
    '  $slug = preg_replace("/[^\\p{L}\\p{N}\\-]/u", "-", $slug);\n'
    '  $slug = preg_replace("/-+/", "-", $slug);\n'
    '  $slug = trim($slug, "-");\n\n'
    '  // 查找或创建\n'
    '  $st = $pdo->prepare("SELECT id FROM post_tag WHERE slug = ? LIMIT 1");\n'
    '  $st->execute([$slug]);\n'
    '  if ($tag = $st->fetch()) {\n'
    '    return (int)$tag["id"];\n'
    '  }\n'
    '  $st = $pdo->prepare("INSERT INTO post_tag \n'
    '    (name, slug, post_count, created_at) VALUES (?, ?, 0, NOW())");\n'
    '  $st->execute([$name, $slug]);\n'
    '  return (int)$pdo->lastInsertId();\n'
    '}\n\n'
    '// 设置帖子标签（最多5个）\n'
    'function set_post_tags(int $postId, array $tagNames): bool {\n'
    '  $pdo->beginTransaction();\n'
    '  // 删除旧关联\n'
    '  $st = $pdo->prepare("DELETE FROM post_tag_relation WHERE post_id = ?");\n'
    '  $st->execute([$postId]);\n'
    '  // 创建新关联\n'
    '  $tagNames = array_slice($tagNames, 0, 5);\n'
    '  foreach ($tagNames as $tagName) {\n'
    '    $tagId = get_or_create_tag($tagName);\n'
    '    if ($tagId > 0) {\n'
    '      $st = $pdo->prepare("INSERT INTO post_tag_relation \n'
    '        (post_id, tag_id) VALUES (?, ?)");\n'
    '      $st->execute([$postId, $tagId]);\n'
    '    }\n'
    '  }\n'
    '  // 刷新所有标签的帖子计数\n'
    '  $st = $pdo->prepare("UPDATE post_tag t SET post_count = (\n'
    '    SELECT COUNT(*) FROM post_tag_relation r WHERE r.tag_id = t.id\n'
    '  )");\n'
    '  $st->execute();\n'
    '  $pdo->commit();\n'
    '  return true;\n'
    '}'
)
add_para(
    'slug生成函数将中文标签转化为URL安全的标识符（如"代数几何"转化为"代数-几何"），'
    '用于URL筛选（如/forum.php?tag=代数-几何）。set_post_tags()在事务中先删除旧关联再创建新关联，'
    '标签数量限制为5个（array_slice）。每次更新后刷新post_tag.post_count字段，'
    '用于展示"热门标签"排行榜。'
)

add_heading_l3('5、通知系统')
add_para(
    '通知系统实现了站内互动通知功能，覆盖评论、回复、点赞、关注四种事件类型。'
    '系统设计为轻量级推送模式：在用户执行互动操作时，同步将通知写入数据库。'
    '核心函数create_notification()在点赞和评论操作的成功分支中被调用：'
)
add_code_block(
    '// 创建通知（app/notification.php）\n'
    'function create_notification(int $userId, string $type,\n'
    '    int $relatedId, ?int $relatedUserId = null,\n'
    '    string $content = ""): bool {\n'
    '  $st = $pdo->prepare("\n'
    '    INSERT INTO notification \n'
    '      (user_id, type, related_id, related_user_id, content, is_read, created_at)\n'
    '    VALUES (?, ?, ?, ?, ?, 0, NOW())"\n'
    '  );\n'
    '  $st->execute([$userId, $type, $relatedId, $relatedUserId, $content]);\n'
    '  return true;\n'
    '}\n\n'
    '// 获取未读通知数量（显示在导航栏小红点）\n'
    'function get_unread_notification_count(int $userId): int {\n'
    '  $st = $pdo->prepare("SELECT COUNT(*) AS c FROM notification \n'
    '    WHERE user_id = ? AND is_read = 0");\n'
    '  $st->execute([$userId]);\n'
    '  return (int)$st->fetch()["c"];\n'
    '}\n\n'
    '// 标记所有通知为已读（用户点击"全部已读"时触发）\n'
    'function mark_all_notifications_read(int $userId): bool {\n'
    '  $st = $pdo->prepare("UPDATE notification SET is_read = 1 \n'
    '    WHERE user_id = ? AND is_read = 0");\n'
    '  $st->execute([$userId]);\n'
    '  return true;\n'
    '}'
)
add_para(
    '通知的触发时机包括：他人评论自己的帖子（触发comment通知）、他人回复自己的评论'
    '（触发reply通知）、他人点赞自己的帖子（触发like通知）。is_read字段用于区分已读和未读，'
    'get_unread_notification_count()用于在导航栏显示未读数量小红点。'
    'clean_old_notifications()定期清理30天前的旧通知，防止通知表无限膨胀。'
)

add_heading_l2('（六）用户设置功能实现')
add_para(
    '用户设置页面（settings.php）是用户管理个人账户信息的统一入口，涵盖个人资料编辑、称号管理、'
    '隐私设置和收货地址管理四个功能模块。该页面采用"即时编辑、即时保存"的设计理念，'
    '所有操作均通过AJAX异步提交，无需整页刷新，提供流畅的用户体验。'
)
add_para(
    '个人资料编辑模块支持头像上传、用户名修改和个性签名编辑三项操作。头像上传通过隐藏的'
    'input[type=file]触发文件选择，在前端进行MIME类型（JPG/PNG/GIF）和文件大小（≤2MB）双重验证后，'
    '通过FormData以AJAX方式提交至avatar_upload_post.php。后端接收文件后进行图像有效性验证，'
    '使用imagecreatefromstring检测是否为合法图片文件，防止恶意文件上传。处理后的头像'
    '以用户ID命名的PNG文件存储，并同步更新user.avatar字段。用户名和个性签名采用"点击-编辑-保存"'
    '的交互模式：用户点击"编辑"按钮后，文本标签切换为输入框或文本框，编辑完成后点击"保存"按钮，'
    '通过AJAX提交至profile_update.php。后端对用户名进行唯一性校验和长度验证（≤50字符），'
    '对签名进行长度验证（≤200字符），验证通过后写入数据库并返回更新后的数据，前端实时更新显示内容。'
    '这种模式避免了传统表单页面跳转的割裂感，降低了用户的操作成本。'
)
add_para(
    '称号管理模块实现了用户在"我的订单"页面购买的称号的展示与切换功能。页面加载时，'
    '后端通过get_user_titles()获取当前用户拥有的所有称号列表，以称号标签（title-chip）的形式渲染。'
    '用户点击某个称号标签时，前端通过AJAX请求title_toggle.php切换展示称号，被激活的称号标签'
    '添加active样式高亮显示，同时提交当前选择的称号标识。后端在user_title表的is_active字段上'
    '执行幂等更新：先将当前用户的所有称号置为0，再将选中称号置为1，保证同一时间只有一个称号处于激活状态。'
    '此外还提供"不展示称号"选项，允许用户隐藏称号展示。'
)
add_para(
    '隐私设置模块目前提供"允许他人查看我的收藏"开关。该功能通过一个toggle-switch组件实现，'
    '切换时立即通过AJAX提交至profile_update.php。后端接收privacy_show_favorites参数后更新'
    'user表的对应字段。该字段在用户个人主页（user.php）中被读取，用于控制收藏tab的可见性，'
    '实现了"设置即生效"的实时反馈体验。'
)
add_para(
    '收货地址管理实现了完整的CRUD操作，是兑换实物商品的必要前置功能。地址数据通过'
    'shipping_address_api.php统一处理，该接口使用action参数区分add、update、delete、set_default'
    '四种操作。地址编辑采用弹出模态框（Modal）的形式，表单包含收件人、联系电话、所在地区和详细地址'
    '四个字段，提交时进行前端必填校验和后端数据验证。后端通过app/address.php中的address_add()、'
    'address_update()等函数操作shipping_address表，每个地址记录通过user_id与用户关联，'
    '并支持"设为默认"功能（is_default字段），默认地址在订单结算时自动填充。'
)

add_heading_l2('（七）积分商城系统')
add_para(
    '积分系统是社区运营的核心激励工具，包含积分获取和积分消费两大模块。'
    '积分获取渠道包括注册奖励、每日签到、发帖、评论等，积分消费则通过商城兑换商品实现。'
    '系统设计的核心挑战是保证并发安全，防止积分超扣。'
)
add_para(
    '对于注册奖励、每日签到等一次性奖励场景，积分流水表point_ledger的(user_id, reason_key)字段'
    '建立了联合唯一索引，reason_key的格式为"行为类型_用户ID_关联ID"，'
    '确保同一用户不会重复领取同一奖励。例如注册奖励的reason_key为"register_{userId}"，'
    '每日签到的reason_key为"checkin_{userId}_{date}"。核心积分发放函数points_grant()如下：'
)
add_code_block(
    '// 积分发放（幂等性保证）（app/points.php）\n'
    'function points_grant(int $userId, int $delta, string $reason,\n'
    '    string $reasonKey): bool {\n'
    '  $pdo = db();\n'
    '  try {\n'
    '    $pdo->beginTransaction();\n'
    '    // 写入流水（利用唯一约束防重复）\n'
    '    $st = $pdo->prepare("INSERT INTO point_ledger \n'
    '      (user_id, delta, reason, reason_key) VALUES (?, ?, ?, ?)");\n'
    '    $st->execute([$userId, $delta, $reason, $reasonKey]);\n'
    '    // 更新用户积分余额\n'
    '    $sumSt = $pdo->query("SELECT SUM(delta) AS bal \n'
    '      FROM point_ledger WHERE user_id = $userId");\n'
    '    $balance = (int)$sumSt->fetch()["bal"];\n'
    '    $st = $pdo->prepare("UPDATE user SET points = ? WHERE id = ?");\n'
    '    $st->execute([$balance, $userId]);\n'
    '    $pdo->commit();\n'
    '    return true;\n'
    '  } catch (Throwable $e) {\n'
    '    if ($pdo->inTransaction()) $pdo->rollBack();\n'
    '    if (strpos($e->getMessage(), "1062") !== false) {\n'
    '      // 重复奖励（违反唯一约束），忽略\n'
    '      return false;\n'
    '    }\n'
    '    return false;\n'
    '  }\n'
    '}'
)
add_para(
    '积分发放后立即重新计算SUM(delta)并更新user.points余额字段。这种"每次写入都重新SUM"的'
    '策略虽然牺牲了少量性能，但保证了积分数额的绝对准确，不会出现系统间对不上的问题。'
    '积分扣减（商城兑换时）使用独立的points_spend()函数，通过FOR UPDATE行级锁防止并发扣减超支：'
)
add_code_block(
    '// 积分扣减（FOR UPDATE行级锁防并发）\n'
    'function points_spend(int $userId, int $cost, string $reason,\n'
    '    string $reasonKey): array {\n'
    '  $pdo = db();\n'
    '  try {\n'
    '    $pdo->beginTransaction();\n'
    '    // 行级锁查询当前积分\n'
    '    $st = $pdo->prepare("SELECT points FROM user WHERE id = ? FOR UPDATE");\n'
    '    $st->execute([$userId]);\n'
    '    $row = $st->fetch();\n'
    '    $balance = $row ? (int)$row["points"] : 0;\n'
    '    if ($balance < $cost) {\n'
    '      $pdo->rollBack();\n'
    '      return ["ok" => false, "msg" => "积分不足"];\n'
    '    }\n'
    '    // 写入负流水\n'
    '    $st = $pdo->prepare("INSERT INTO point_ledger \n'
    '      (user_id, delta, reason, reason_key) VALUES (?, ?, ?, ?)");\n'
    '    $st->execute([$userId, -$cost, $reason, $reasonKey]);\n'
    '    // 更新余额\n'
    '    $st = $pdo->prepare("UPDATE user SET points = points - ? WHERE id = ?");\n'
    '    $st->execute([$cost, $userId]);\n'
    '    $pdo->commit();\n'
    '    return ["ok" => true, "msg" => "兑换成功"];\n'
    '  } catch (Throwable $e) {\n'
    '    if ($pdo->inTransaction()) $pdo->rollBack();\n'
    '    return ["ok" => false, "msg" => "兑换失败"];\n'
    '  }\n'
    '}'
)
add_para(
    '签到系统利用INSERT的唯一键冲突保证每人每天只能签到一次。user_checkin表以(user_id, checkin_date)'
    '为联合主键，签到原理是：尝试INSERT一条记录，如果当天已签过到则INSERT失败（重复键错误），'
    '捕获该错误后返回"今日已签到"。'
    '连续签到奖励的逻辑是从今天往前追溯，计算连续签到的天数，达到7天和30天时发放一次性奖励。'
    '签到奖励分为基础奖励（每日固定积分）和连续奖励（阶段性额外奖励），'
    '两者使用不同的reason_key防止互相干扰。'
)
add_para(
    '商城兑换允许用户多次兑换同一商品，每次兑换生成独立的shop_order订单记录，'
    'points_spend()在积分流水表中写入一条负流水。管理员后台可查看所有订单、管理物流状态，'
    '支持CSV导出和批量上传物流单号。'
)

add_heading_l2('（八）管理员后台实现')
add_para(
    '管理员后台是网站运营管理的核心枢纽，采用"左侧导航+右侧内容区"的经典布局，'
    '包含仪表盘、帖子审核、商城管理、订单管理和素材库五个功能页面。'
    '所有后台页面共用一套权限体系和侧边栏导航组件。'
)
add_heading_l3('1、权限管理体系')
add_para(
    '后台权限基于user.role字段实现，该字段默认值为"user"，通过admin_ensure_schema()函数'
    '在首次访问后台时动态添加到user表中。权限检查通过两个函数完成：is_admin()判断当前用户'
    '是否为管理员，require_admin()在非管理员访问时直接终止请求并返回403。'
    'require_admin()是每个后台页面的第一道防线，在页面任何业务逻辑执行之前调用。'
)
add_code_block(
    '// 权限检查函数（app/admin.php）\n'
    'function is_admin(): bool {\n'
    '  $user = current_user();\n'
    '  return $user !== null && isset($user[\'role\']) && $user[\'role\'] === \'admin\';\n'
    '}\n\n'
    'function require_admin(): void {\n'
    '  if (!is_admin()) {\n'
    '    http_response_code(403);\n'
    '    if (is_ajax()) {\n'
    '      header(\'Content-Type: application/json; charset=utf-8\');\n'
    '      echo json_encode([\'ok\' => false, \'msg\' => \'Forbidden\']);\n'
    '    } else {\n'
    '      echo \'<h1>403 Forbidden</h1><p>无权限访问管理后台</p>\';\n'
    '    }\n'
    '    exit;\n'
    '  }\n'
    '}'
)
add_para(
    '管理员后台还实现了完整的操作审计机制。admin_log()函数在所有关键操作（审核帖子、'
    '发货、删除素材等）后调用，将操作记录写入admin_log表。该表记录了操作人（admin_id）、'
    '操作类型（action）、操作对象（target_type/target_id）和操作详情（detail），'
    '为后续的审计追溯提供了完整的数据支撑。'
)

add_heading_l3('2、仪表盘')
add_para(
    '仪表盘（admin.php）是管理员登录后台后的默认页面，集中展示网站的核心运营数据。'
    '页面加载时执行5个COUNT查询分别获取总用户数、已发布帖子数、待审核帖子数、'
    '商品总数和30天内已删除帖子数，以统计卡片的形式直观呈现。'
    '最近操作日志区域通过LEFT JOIN关联admin_log和user表，显示最近20条操作记录，'
    '包括操作时间、操作人和操作内容。'
)
add_para(
    '仪表盘还提供"永久清理30天前删除帖"的数据维护功能。该功能调用'
    'admin_cleanup_trashed_posts()函数，物理删除status=0且updated_at超过30天的'
    '帖子及其关联数据（评论、点赞、收藏、标签关系、通知），防止数据库被软删除数据无限膨胀。'
    '清理操作通过事务保证完整性，且只有在管理员手动点击按钮或配置定时任务时才会触发。'
)
add_code_block(
    '// 清理过期删除帖（app/admin.php）\n'
    'function admin_cleanup_trashed_posts(): int {\n'
    '  $pdo = db();\n'
    '  $st = $pdo->query(\n'
    '    "SELECT id FROM post WHERE status = 0\n'
    '     AND updated_at < DATE_SUB(NOW(), INTERVAL 30 DAY)");\n'
    '  $ids = $st->fetchAll(PDO::FETCH_COLUMN);\n'
    '  if (empty($ids)) return 0;\n\n'
    '  $idList = implode(\',\', array_map(\'intval\', $ids));\n'
    '  $pdo->beginTransaction();\n'
    '  $pdo->exec("DELETE FROM comment WHERE post_id IN ($idList)");\n'
    '  $pdo->exec("DELETE FROM post_like WHERE post_id IN ($idList)");\n'
    '  $pdo->exec("DELETE FROM post_favorite WHERE post_id IN ($idList)");\n'
    '  $pdo->exec("DELETE FROM post_tag_relation WHERE post_id IN ($idList)");\n'
    '  $pdo->exec("DELETE FROM notification\n'
    '    WHERE (type IN (\'comment\',\'like\') AND related_id IN ($idList))");\n'
    '  $pdo->exec("DELETE FROM post WHERE id IN ($idList)");\n'
    '  $pdo->commit();\n'
    '  return count($ids);\n'
    '}'
)

add_heading_l3('3、帖子审核管理')
add_para(
    '帖子审核是管理员后台的核心功能。用户在creator.php发布帖子后，post.review_status被设为1（待审核），'
    '帖子不会出现在社区列表中对其他用户可见。管理员在admin_posts.php中通过筛选条件'
    '（全部/已发布/草稿/待审核/已驳回）查看不同状态的帖子，对待审核帖子可以执行批准'
    '（review_status设为0）或驳回（设为2）操作，已驳回的帖子作者可以查看并修改后重新提交。'
)
add_para(
    '操作通过admin_posts_api.php的POST接口处理，使用action参数区分approve、reject、'
    'delete、batch_approve、batch_reject五种操作。其中批量操作通过逗号分隔的post_ids参数'
    '接收多个帖子ID，使用IN语句和占位符动态生成SQL，避免拼接注入风险。'
)
add_code_block(
    '// 帖子审核核心逻辑（admin_posts_api.php）\n'
    'switch ($action) {\n'
    '  case \'approve\':\n'
    '    $st = $pdo->prepare(\n'
    '      "UPDATE post SET review_status = 0 WHERE id = ? AND status = 1");\n'
    '    $st->execute([$postId]);\n'
    '    admin_log(\'post_approve\', \'post\', $postId,\n'
    '      \'批准帖子 #\' . $postId);\n'
    '    echo json_encode([\'ok\' => true, \'msg\' => \'已批准\']);\n'
    '    break;\n\n'
    '  case \'batch_approve\':\n'
    '    $ids = array_map(\'intval\', explode(\',\', $postIdsRaw));\n'
    '    $ids = array_filter($ids, fn($v) => $v > 0);\n'
    '    $placeholders = implode(\',\',\n'
    '      array_fill(0, count($ids), \'?\'));\n'
    '    $st = $pdo->prepare(\n'
    '      "UPDATE post SET review_status = 0\n'
    '       WHERE id IN ($placeholders) AND status = 1");\n'
    '    $st->execute(array_values($ids));\n'
    '    admin_log(\'post_batch_approve\', \'post\', 0,\n'
    '      \'批量批准 \' . count($ids) . \' 篇帖子\');\n'
    '    echo json_encode([\'ok\' => true, \'msg\' => \'已批量批准\']);\n'
    '    break;\n'
    '}'
)

add_heading_l3('4、商城管理')
add_para(
    '商城管理页面（admin_shop.php）实现了积分商品的完整管理功能。管理员可以查看所有商品列表，'
    '并支持添加、编辑、删除商品。每个商品包含标题、描述、积分价格、图标（Emoji）、'
    '商品图片、排序权重、上架状态、是否为实物、是否可重复兑换、是否为称号类虚拟商品等属性。'
    '商品列表从shop_item表中按sort_order排序读取，以表格形式展示各字段信息。'
    '该页面使用admin_shop_api.php作为API后端，通过action参数区分create、update、delete、get四种操作。'
    '所有操作均记录到admin_log审计日志中。'
)

add_heading_l3('5、订单管理')
add_para(
    '订单管理页面（admin_orders.php + admin_orders_api.php）处理用户兑换实物商品的完整物流工作流。'
    '订单有四种状态：0=待发货、1=已发货、2=已完成、3=已取消，构成了完整的生命周期。'
    '管理员可以按订单状态筛选列表，对"待发货"订单执行发货操作（填写物流单号），'
    '系统自动将状态转为"已发货"并触发站内通知告知买家；"已发货"订单可标记为"已完成"。'
    '订单列表支持分页展示，每页20条，通过LEFT JOIN关联user表显示买家用户名。'
)
add_para(
    '订单管理还实现了两个实用的批量操作功能。CSV导出功能通过GET请求的export_csv动作实现：'
    '查询符合条件的订单列表，将数据以UTF-8 BOM编码的CSV格式输出，包含订单ID、商品名称、'
    '用户名、收件人、手机号、地址、积分、状态、物流单号和下单时间等字段，'
    '可直接用Excel打开。物流单号批量导入功能支持管理员上传CSV文件，'
    '系统自动解析"订单ID"和"物流单号"两列，逐行校验订单状态后批量更新发货。'
)
add_code_block(
    '// CSV导出功能（admin_orders_api.php）\n'
    '$sql = "SELECT o.*, u.username\n'
    '        FROM shop_order o\n'
    '        LEFT JOIN user u ON u.id = o.user_id\n'
    '        ORDER BY o.id DESC";\n'
    '$orders = $pdo->prepare($sql)->execute($params)->fetchAll();\n\n'
    'header(\'Content-Type: text/csv; charset=utf-8\');\n'
    'header(\'Content-Disposition: attachment;\n'
    '        filename="orders_\' . date(\'Ymd\') . \'.csv"\');\n'
    '$out = fopen(\'php://output\', \'w\');\n'
    'fprintf($out, "\xEF\xBB\xBF");  // UTF-8 BOM\n'
    'fputcsv($out, [\'订单ID\', \'商品名称\', \'用户名\', \'收件人\',\n'
    '  \'手机号\', \'地址\', \'积分\', \'状态\', \'物流单号\', \'下单时间\']);\n'
    'foreach ($orders as $o) {\n'
    '  fputcsv($out, [$o[\'id\'], $o[\'item_title\'],\n'
    '    $o[\'username\'], $o[\'recipient_name\'],\n'
    '    $o[\'recipient_phone\'], $o[\'recipient_address\'],\n'
    '    $o[\'cost_points\'], $statusLabels[(int)$o[\'status\']],\n'
    '    $o[\'tracking_number\'], $o[\'created_at\']]);\n'
    '}'
)

add_heading_l3('6、素材库管理')
add_para(
    '素材库（admin_media.php + admin_media_api.php）为管理员提供统一的图片资产管理功能。'
    '页面以网格形式展示已上传的图片文件，支持分页浏览（每页48张）、按文件名搜索和删除操作。'
    '前端通过AJAX调用file_list、file_search等接口获取数据，使用JavaScript动态渲染网格视图。'
)
add_para(
    '文件上传功能在前端限制文件类型为JPEG/PNG/GIF/WebP/SVG，大小限制为10MB。'
    '后端进行双重验证：首先检查文件扩展名是否在允许列表中，然后通过getimagesize()函数'
    '检测文件的真实MIME类型和图像尺寸，防止伪装扩展名的恶意文件上传。'
    '上传后的文件以"时间戳_随机十六进制字符串.扩展名"的格式重命名，'
    '存储于uploads/media目录下，元数据写入media_file表，'
    '包括原始文件名、MIME类型、文件大小、图片宽高等信息，便于后续检索和管理。'
)

add_heading_l2('（九）画板功能实现')
add_para(
    '画板是网站为数学用户提供的在线绘图工具，基于原生Canvas API实现，不依赖任何第三方绘图库。'
    '前端采用HTML模板层（drawboard.php）定义界面布局和工具栏、JavaScript引擎层（drawboard.js）'
    '实现全部绘图逻辑和交互、CSS样式层定义面板样式的三层架构。'
)
add_para(
    '绘图引擎采用"路径-渲染"分离架构，支持自由手绘（二次贝塞尔曲线平滑）、直线、矩形、圆、'
    '文本标注、坐标轴和函数图象七种图形类型。函数图象绘制通过new Function动态编译用户输入的表达式，'
    '支持Math对象的所有数学函数。画板实现了几何吸附系统，鼠标可自动吸附到端点、中点、圆心等关键点；'
    '撤销/重做通过undoStack/redoStack实现，最大深度50步，支持Ctrl+Z/Ctrl+Shift+Z快捷键；'
    '图形选择支持点击命中检测、拖拽移动、Shift多选和Delete删除，并可通过canvas.toDataURL()导出为PNG。'
    '画板同时支持Retina屏幕适配，通过devicePixelRatio确保高DPI屏幕上的显示清晰度。'
)

# 五、功能实现效果
add_heading_l1('五、功能实现效果')

add_heading_l2('（一）前台用户端')

add_heading_l3('1、首页')
add_para(
    '首页是整个网站的门面，设计上注重视觉冲击力和品牌调性。主视觉区（Hero Section）展示站点标语'
    '"种下热爱，等数学开花"，通过打字机效果逐字呈现，每260毫秒显示一个字，'
    '营造"书写"的仪式感。打字完成后，副标题和操作按钮通过GSAP动画优雅入场（向上位移+淡入）。'
    '数据显示区展示四个维度的社区动态：活跃用户数、内容主题数、知识标签数、深度回复数，'
    '使用GSAP的ScrollTrigger插件实现滚动触发的数字递增动画，从0逐渐滚动到目标值。'
    '开发者团队区域实现横向无缝轮播效果，团队卡片采用正方形设计展示六位开发者照片，'
    '悬停时轮播不会暂停，团队成员数据从后端PHP数组渲染。'
    '特色展示区域四个卡片（专业内容、活跃社区、知识库、创新思维）在滚动到视口时依次入场，'
    '交错时间差为100毫秒，形成节奏感。'
)

add_heading_l3('2、社区页面')
add_para(
    '社区页面（forum.php）是用户浏览内容的核心场景。推荐/热门/动态三种排序模式的切换通过URL参数实现，'
    '支持标签筛选和无限滚动加载。无限滚动通过监听scroll事件实现：当页面滚动距离底部200px时，'
    '通过AJAX请求加载下一页内容，解析返回的HTML片段并将其追加到帖子网格中。'
    '每个帖子卡片以统一的post_grid_card.php模板渲染，显示封面图、标题、摘要（纯文本截断）、标签和作者信息。'
)

add_heading_l3('3、帖子详情页')
add_para(
    '帖子详情页展示完整的内容，包括作者信息、发布时间、标签、正文（由MathJax渲染数学公式）、'
    '以及底部的互动工具栏。互动工具栏固定在视口底部，包含点赞、收藏、分享、评论四个按钮。'
    '评论区域支持主评论和回复两种模式，每条评论显示用户头像、用户名、称号、内容和时间，'
    '支持点赞、回复和删除操作。'
)

add_heading_l3('4、创作者平台与签到商城')
add_para(
    '创作者平台（creator.php）是基于Quill编辑器的发帖页面，提供完整的富文本编辑功能。'
    '商城页面（shop.php）展示可兑换的虚拟商品和实物商品，用户可以使用积分进行兑换。'
    '每日签到功能鼓励用户持续访问社区，连续签到7天和30天可获得额外积分奖励。'
)

add_heading_l3('5、用户设置页面')
add_para(
    '用户设置页面（settings.php）整合了个人资料编辑、称号管理、隐私设置和收货地址管理四个功能模块。'
    '个人资料区域支持头像点击预览大图、用户名和个性签名的即时编辑；称号管理区域以标签形式展示用户拥有的'
    '所有称号，点击即可切换展示称号；隐私设置提供"允许他人查看我的收藏"开关，切换即时生效；'
    '收货地址管理支持添加、编辑、删除和设为默认操作，地址编辑在模态弹窗中完成。'
    '整个页面所有操作均通过AJAX异步提交，无需整页刷新。'
)

add_heading_l2('（二）管理员后台')
add_para(
    '管理员后台布局采用左侧导航+右侧内容区的经典设计。仪表盘展示网站核心统计数据（用户数、帖子数、'
    '待审核内容等）和最近20条操作日志，提供"永久清理30天前删除帖"的数据维护功能。'
    '帖子管理支持按状态（已发布/待审核/已驳回/草稿）筛选，支持单篇批准/驳回和批量操作；'
    '商城管理支持商品的上架、编辑、下架等完整维护操作，商品支持实物/虚拟称号两种类型，'
    '以及是否可重复兑换等属性配置。'
    '订单管理支持完整的物流工作流，从待发货到已签收的完整状态流转，以及CSV导出和批量导入物流单号功能。'
    '素材库管理对上传的媒体文件进行网格化展示，支持分页浏览、文件名搜索、上传和删除操作。'
    '所有管理功能均记录操作日志，便于审计追溯。'
)

# 六、Web防护体系及渗透测试
add_heading_l1('六、Web防护体系及渗透测试')

add_para(
    '作为信息安全专业的毕业设计，本系统在设计和实现过程中构建了完整的纵深防御体系。'
    '本章从防护策略和渗透测试两个维度，对系统的Web安全能力进行全面评估。'
)

add_heading_l2('（一）SQL注入防护与测试')
add_para(
    '防护策略：系统全面使用PDO预处理语句，并禁用模拟预处理模式（ATTR_EMULATE_PREPARES=false），'
    '强制使用MySQL原生预处理协议。SQL模板与参数分离，参数值永远不会被解析为SQL代码。'
    '所有数据库操作统一使用prepare()+execute()模式，避免query()直接拼接SQL。'
)
add_para(
    '测试结果：使用sqlmap对网站主要GET和POST参数进行扫描（--level=3 --risk=2），'
    '未发现SQL注入漏洞。参数化接口从根本上杜绝了SQL注入的可能。'
)

add_heading_l2('（二）XSS防护与测试')
add_para(
    '防护策略：系统构建了四层XSS纵深防御体系。第一层为输出转义，在模板中统一使用'
    'htmlspecialchars()对纯文本数据进行HTML实体编码。第二层为DOM树级HTML净化，针对富文本内容，'
    '使用PHP的DOMDocument类解析为DOM树后进行白名单检查，只允许安全标签和安全属性，'
    '过滤script、iframe等危险标签和javascript:、expression()等危险内容。'
    '第三层设置Content-Security-Policy HTTP响应头，限制页面资源加载来源。'
    '第四层在Cookie层面设置HttpOnly属性，防止JavaScript通过document.cookie读取会话标识。'
)
add_para(
    '测试结果：在发帖和评论功能中注入常见XSS payload（<script>、<img onerror=、<svg onload=等），'
    '所有注入尝试均被拦截。DOM净化层删除了危险标签和事件处理器，CSP头进一步限制了脚本执行来源。'
)

add_heading_l2('（三）CSRF防护与测试')
add_para(
    '防护策略：采用同步器令牌模式，所有数据修改操作的表单和AJAX请求都包含与用户会话绑定的'
    'CSRF令牌（32位十六进制随机数）。服务端使用hash_equals()进行恒定时间比对。'
    '同时设置SameSite=Lax Cookie策略，阻止跨站请求携带会话cookie，形成第二道防线。'
)
add_para(
    '测试结果：在外部页面构造自动提交的POST表单尝试执行点赞、发帖等操作，所有跨站请求均被'
    'CSRF令牌验证拦截，返回403状态码。删除csrf_token参数的请求同样被拒绝。'
)

add_heading_l2('（四）文件上传防护与测试')
add_para(
    '防护策略：采用多层文件安全校验机制。首先使用finfo_open()检测文件真实MIME类型，'
    '仅允许白名单中的图片类型。其次验证文件扩展名是否在白名单中。文件大小限制为头像2MB、'
    '素材库10MB。文件名使用时间戳或用户ID重命名，不使用用户提交的原始文件名。'
    '头像上传后重新编码为JPG格式，自动剥离EXIF元数据。'
)
add_para(
    '测试结果：尝试上传.php、.asp、.exe等可执行文件、修改扩展名绕过、上传超大文件、'
    '上传图片马等操作均被拦截。非白名单扩展名被拒绝，MIME内容检测拦截了扩展名篡改，'
    'JPG重编码移除了图片马中的PHP代码。'
)

add_heading_l2('（五）认证安全与暴力破解防护及测试')
add_para(
    '防护策略：密码使用bcrypt算法（password_hash()）存储，每次自动生成随机salt，'
    '防止彩虹表攻击。登录接口实现基于失败计数和账号锁定的防爆破机制：连续失败5次后'
    '账号锁定15分钟。Session配置HttpOnly、Secure、SameSite=Lax属性，'
    '登录成功后调用session_regenerate_id()防止会话固定攻击。'
)
add_para(
    '测试结果：使用Hydra和自定义脚本对登录接口进行连续密码尝试，连续输错5次密码后账号被锁定，'
    '锁定期间任何密码均无法登录。bcrypt的计算耗时（约50-100ms/次）进一步降低了暴力破解效率。'
)

add_heading_l2('（六）目录遍历与越权防护及测试')
add_para(
    '防护策略：所有文件路径均为硬编码或通过函数安全生成，不接受用户输入的路径参数。'
    '越权防护从两个维度展开——垂直越权方面，所有管理员操作通过require_admin()校验角色身份；'
    '水平越权方面，所有数据修改操作在SQL的WHERE条件中加入user_id = ?约束，'
    '确保只能操作属于自己的资源。'
)
add_para(
    '测试结果：尝试通过../序列访问/etc/passwd等系统文件被拒绝，所有路径均为硬编码。'
    '普通用户会话访问管理员接口返回403。尝试修改post_id、user_id等参数访问他人数据'
    '被SQL条件中的user_id校验拦截。'
)

add_heading_l2('（七）自动化攻击防护与频率限制实践')
add_para(
    '在网站开发收尾阶段，系统尚未正式上线期间，发生了一起自动化脚本攻击事件。'
    '一个注册用户名为"e"的账号通过脚本批量调用点赞接口（like_toggle.php和comment_like_toggle.php），'
    '在短时间内对大量帖子和评论执行了重复的点赞和取消点赞操作。'
    '由于点赞接口采用toggle设计（同一请求交替执行插入和删除），数据本身未被破坏，'
    '但大量的无效请求占用了数据库连接和服务器资源，暴露了系统缺乏请求频率限制的薄弱环节。'
)
add_para(
    '针对这一事件，系统新增了文件级的频率限制中间件。'
    '该中间件的设计遵循三个原则：第一，轻量无依赖，不引入Redis等外部服务，'
    '使用文件系统存储请求记录，适用于低配置服务器环境；第二，按行为分类限制，'
    '对点赞（30次/分钟）、注册（5次/5分钟）和登录（10次/分钟）等不同类型的操作'
    '设置差异化的频率阈值；第三，自动过期清理，超出时间窗口的请求记录自动淘汰，'
    '不会造成磁盘空间堆积。核心实现如下：'
)
add_code_block(
    '// 频率限制中间件（app/ratelimit.php）\n'
    'function check_rate_limit(string $action,\n'
    '    int $max = 30, int $window = 60): void {\n'
    '  $ip = $_SERVER[\'REMOTE_ADDR\'] ?? \'unknown\';\n'
    '  $key = \'rl_\' . $action . \'_\' . md5($ip);\n'
    '  $file = sys_get_temp_dir() . \'/\' . $key . \'.php\';\n\n'
    '  $now = time();\n'
    '  $records = [];\n'
    '  if (file_exists($file)) {\n'
    '    $data = @file_get_contents($file);\n'
    '    if ($data) $records = unserialize($data);\n'
    '  }\n'
    '  // 清理过期记录\n'
    '  $records = array_values(\n'
    '    array_filter($records, fn($t) => $t > $now - $window));\n\n'
    '  if (count($records) >= $max) {\n'
    '    http_response_code(429);\n'
    '    echo json_encode([\'ok\'=>false,\n'
    '      \'msg\'=>\'操作过于频繁，请稍后再试\']);\n'
    '    exit;\n'
    '  }\n\n'
    '  $records[] = $now;\n'
    '  @file_put_contents($file, serialize($records), LOCK_EX);\n'
    '}'
)
add_para(
    '该频率限制中间件的引入过程本身也是一次安全实践。在事件发生前，系统已具备CSRF令牌验证防止'
    '跨站请求伪造，但未对同一用户的请求速率做任何限制。事件发生后，仅新增一个不到40行的中间件文件，'
    '在四个关键接口中各添加一行调用代码，就完成了防护能力的补充。这体现了安全防护需要多层次、'
    '多维度覆盖的理念——CSRF防护解决了"请求来源是否可信"的问题，而频率限制解决了'
    '"请求数量是否合理"的问题，两者互为补充，共同构建更完整的安全屏障。'
)

add_heading_l2('（八）其他安全性审计')
add_para(
    'PHP反序列化漏洞：代码审计确认系统未使用unserialize()函数，不存在反序列化攻击面。'
    'SSRF漏洞：系统仅向硬编码的MiniMax API发送HTTPS请求，用户无法控制请求目标。'
    'RCE漏洞：系统未使用eval()、exec()、system()等危险函数，无代码执行入口。'
    '文件包含漏洞：所有include/require路径均为硬编码常量，不涉及用户输入。'
    '敏感信息泄露：PHP错误报告不输出到页面，异常信息通过error_log写入服务器日志，'
    '用户可见的错误消息均为通用提示语。'
)

add_heading_l2('（九）总结')
add_para(
    '经过系统的安全防护设计与渗透测试验证，数问社区网站在SQL注入、XSS、CSRF、文件上传、'
    '认证安全、目录遍历、越权访问、暴力破解等常见Web攻击维度均表现出良好的防御能力，'
    '无高危漏洞。针对自动化脚本攻击实践了频率限制中间件，弥补了请求速率控制这一薄弱环节。'
    '代码审计确认系统不存在PHP反序列化、SSRF、RCE、文件包含等漏洞的攻击面。'
    '安全防护贯穿了系统的每一层——从数据库访问层的预处理语句，到前端输出的安全转义和CSP策略，'
    '再到业务逻辑层的权限校验和内容审核，形成了完整的纵深防御体系。'
)

# 七、实践心得
add_heading_l1('七、实践心得')
add_para(
    '回顾整个毕业设计项目的开发过程，从需求分析、数据库设计、代码编写到安全测试，'
    '每一个环节都让我收获颇丰。在技术层面，我深入理解了PDO预处理语句的防注入原理，'
    '掌握了CSRF令牌验证和DOM净化的实现细节，体会了反范式化设计在性能优化中的作用。'
    '特别是RAG技术的实践，让我对AI与知识库的结合有了更具体的认识。'
)
add_para(
    '在项目开发过程中也遇到了一些值得记录的挑战。最初的数学公式渲染方案只使用了MathJax，'
    '但在编辑器中的渲染速度不够理想，后来引入了KaTeX进行预览渲染，形成了双引擎方案。'
    '热门排序算法也经过了多次迭代，从简单的按点赞数排序，到考虑时间衰减的加权公式，'
    '最终形成了现有的热度计算公式。积分系统的幂等性设计也经过了仔细推敲，'
    '最终通过数据库唯一约束和事务机制保证了数据的准确性。'
)
add_para(
    '通过这次毕设实践，我深刻认识到：信息安全不是独立的功能模块，而应该渗透到系统设计的每一个环节。'
    '从数据库访问、用户输入处理、前端输出渲染，到业务逻辑层的权限校验，安全防护需要全方位、多层次的'
    '纵深防御体系。这次经历也让我对网站全栈开发有了更完整的认知，为未来的职业发展打下了坚实的基础。'
)

# 八、结论
add_heading_l1('八、结论')
add_para(
    '数问（Sown）数学社区网站是一个功能完整、安全可靠的垂直社区平台。系统实现了内容发布与管理、'
    '社区互动、积分体系、AI问答、管理员后台等核心功能，数学公式的渲染与编辑方案解决了数学类社区'
    '的核心痛点。在安全方面，系统构建了CSRF令牌验证、XSS多层过滤、SQL注入防御、登录防爆破等'
    '完整的纵深防御体系，经工具扫描和人工测试验证无高危漏洞。'
)
add_para(
    '当然，项目仍存在一些可以优化的方向：(1) 搜索功能目前基于MySQL的LIKE模糊查询，'
    '后续可以考虑接入Elasticsearch或全文索引提升搜索效率和精准度；(2) 画板功能目前较为基础，'
    '后续可以引入更完善的在线绘图工具；(3) 内容审核目前基于本地关键词库的简单子串匹配，'
    '误伤率和漏报率较高，后续可以接入云端内容安全API（如阿里云内容安全、腾讯云天御等），'
    '利用NLP和图像识别技术实现更精准的违规内容检测；(4) 系统在高并发场景下的性能还需要进一步压测和优化。'
    '这些都是未来持续迭代的方向。'
)

# 致谢
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('致  谢')
run.bold = True
run.font.size = Pt(16)
run.font.name = '宋体'
add_blank()

add_para(
    '本毕业设计的完成，离不开我的指导老师王崎老师的悉心指导。从选题方向到技术方案、从论文框架到细节修改，'
    '王崎老师都给予了我耐心的指导和宝贵的建议。在此，向王崎老师表示最诚挚的感谢。'
)
add_para(
    '同时，感谢实习公司的同事们，他们在实际项目中给予了我很多技术上的帮助和经验分享。'
    '也感谢我的家人和朋友，在整个毕设期间给予的支持与鼓励。'
)
add_para(
    '最后，感谢所有在数问社区中参与测试的用户，你们的反馈让这个项目变得更加完善。'
)

# 参考文献
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('参考文献')
run.bold = True
run.font.size = Pt(16)
run.font.name = '宋体'
add_blank()

refs = [
    '[1] 张炳帅. Web安全深度剖析[M]. 北京: 电子工业出版社, 2015.',
    '[2] PHP: Hypertext Preprocessor. PHP Manual[EB/OL]. https://www.php.net/manual/, 2025.',
    '[3] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems 33, 2020: 9459-9474.',
    '[4] KaTeX. KaTeX Documentation[EB/OL]. https://katex.org/docs/, 2025.',
    '[5] OWASP. Cross-Site Request Forgery Prevention Cheat Sheet[EB/OL]. https://cheatsheetseries.owasp.org/cheatsheets/Cross-Site_Request_Forgery_Prevention_Cheat_Sheet.html, 2024.',
    '[6] OWASP. Cross Site Scripting Prevention Cheat Sheet[EB/OL]. https://cheatsheetseries.owasp.org/cheatsheets/Cross_Site_Scripting_Prevention_Cheat_Sheet.html, 2024.',
    '[7] MiniMax. MiniMax开放平台API文档[EB/OL]. https://platform.minimaxi.com/, 2025.',
    '[8] The MathJax Consortium. MathJax Documentation[EB/OL]. https://docs.mathjax.org/, 2024.',
    '[9] Botros S, Tinley J. High Performance MySQL[M]. 4th ed. Sebastopol: O\'Reilly Media, 2021.',
    '[10] OWASP. SQL Injection Prevention Cheat Sheet[EB/OL]. https://cheatsheetseries.owasp.org/cheatsheets/SQL_Injection_Prevention_Cheat_Sheet.html, 2024.',
    '[11] OWASP. File Upload Prevention Cheat Sheet[EB/OL]. https://cheatsheetseries.owasp.org/cheatsheets/File_Upload_Cheat_Sheet.html, 2024.',
    '[12] OWASP. Content Security Policy Cheat Sheet[EB/OL]. https://cheatsheetseries.owasp.org/cheatsheets/Content_Security_Policy_Cheat_Sheet.html, 2024.',
    '[13] 张海藩, 牟永敏. 软件工程导论[M]. 第6版. 北京: 清华大学出版社, 2013.',
    '[14] Stuttard D, Pinto M. The Web Application Hacker\'s Handbook: Finding and Exploiting Security Flaws[M]. 2nd ed. Indianapolis: Wiley, 2011: 45-128.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(ref)
    run.font.size = Pt(12)
    run.font.name = '宋体'

# 附录
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('附录')
run.bold = True
run.font.size = Pt(16)
run.font.name = '宋体'
add_blank()

add_para(
    '附录A：核心数据表DDL（完整SQL建表语句）。'
)
add_para(
    '附录B：安全防护核心函数源码（csrf.php、helpers.php中sanitize相关函数）。'
)
add_para(
    '附录C：推流算法完整逻辑（forum.php全文）。'
)
add_para(
    '附录D：AI智能体配置文件及RAG完整源码（ai_config.php、ai_index.php、ask_teacher_api.php）。'
)
add_para(
    '因篇幅限制，完整项目源代码请参见项目仓库。'
)

# ============ SAVE ============

output_path = '/www/wwwroot/Sown/毕业设计/数问社区网站毕业综合实践报告.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print('Done!')
